#!/usr/bin/env python3
"""
Convert markdown to .docx optimized for Google Docs upload compatibility.

Strategy: minimize fancy features that Google Docs rejects.
  - Only use built-in core styles: Normal + Heading 1-4
  - Lists rendered as paragraphs with "•" prefix (no list-numbering XML)
  - Tables with plain Table Grid style only
  - Images resized to max 1200px wide, converted to JPEG, stripped EXIF
  - Code blocks as monospace paragraphs (no shading XML)
  - Blockquotes as italic paragraphs with indent
"""
from __future__ import annotations

import argparse
import io
import sys
from pathlib import Path
from urllib.parse import unquote

import markdown as md
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

DOCX_HEADING_LEVELS = {
    "h1": 0,
    "h2": 1,
    "h3": 2,
    "h4": 3,
    "h5": 4,
    "h6": 5,
}

MAX_IMG_WIDTH = 1200
MAX_IMG_INCHES = 5.5


def add_inline_runs(paragraph, element):
    for node in element.children if hasattr(element, "children") else [element]:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                paragraph.add_run(text)
            continue
        name = node.name
        if name in ("strong", "b"):
            run = paragraph.add_run(node.get_text())
            run.bold = True
        elif name in ("em", "i"):
            run = paragraph.add_run(node.get_text())
            run.italic = True
        elif name == "code":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif name == "a":
            run = paragraph.add_run(node.get_text())
            run.font.color.rgb = RGBColor(0x00, 0x4C, 0x99)
            run.underline = True
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            add_inline_runs(paragraph, node)


def add_paragraph(doc, element):
    p = doc.add_paragraph()
    add_inline_runs(p, element)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_list_item(doc, element, ordered=False, index=1):
    """Render as plain paragraph with bullet/number prefix. No list-style XML."""
    p = doc.add_paragraph()
    prefix = f"{index}. " if ordered else "•  "
    p.add_run(prefix)
    add_inline_runs(p, element)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_table(doc, table_el):
    rows = table_el.find_all("tr")
    if not rows:
        return
    cols = max(len(tr.find_all(["th", "td"])) for tr in rows)
    if cols == 0:
        return
    docx_table = doc.add_table(rows=len(rows), cols=cols)
    docx_table.style = "Table Grid"
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for c_idx, td in enumerate(cells):
            if c_idx >= cols:
                continue
            cell = docx_table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, td)
            if r_idx == 0 or td.name == "th":
                for run in p.runs:
                    run.bold = True


def add_image(doc, img_el, base_dir):
    src = img_el.get("src", "")
    alt = img_el.get("alt", "")
    src = unquote(src)
    img_path = (base_dir / src).resolve()
    if img_path.exists() and img_path.is_file():
        try:
            img = Image.open(img_path)
            # Convert RGBA → RGB on white background (JPEG doesn't support alpha)
            if img.mode in ("RGBA", "LA", "P"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                bg.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
                img = bg
            elif img.mode != "RGB":
                img = img.convert("RGB")
            # Resize if too large
            if img.size[0] > MAX_IMG_WIDTH:
                ratio = MAX_IMG_WIDTH / img.size[0]
                new_size = (MAX_IMG_WIDTH, int(img.size[1] * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            # Save as JPEG in memory
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85, optimize=True)
            buf.seek(0)
            doc.add_picture(buf, width=Inches(MAX_IMG_INCHES))
            if alt:
                caption_p = doc.add_paragraph()
                caption_run = caption_p.add_run(alt)
                caption_run.italic = True
                caption_run.font.size = Pt(9)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        except Exception as e:
            print(f"  WARN: nao consegui inserir imagem {img_path}: {e}", file=sys.stderr)
    p = doc.add_paragraph()
    run = p.add_run(f"[Imagem: {alt or src}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_blockquote(doc, element):
    """Render as italic paragraph with left indent."""
    p = doc.add_paragraph()
    for child in element.children:
        if child.name == "p":
            for sub in child.children:
                if isinstance(sub, NavigableString):
                    run = p.add_run(str(sub))
                    run.italic = True
                else:
                    run = p.add_run(sub.get_text())
                    run.italic = True
        else:
            if isinstance(child, NavigableString):
                run = p.add_run(str(child))
                run.italic = True
    p.paragraph_format.left_indent = Inches(0.4)


def process_element(doc, element, base_dir):
    """Walk a single top-level element and write to doc."""
    if isinstance(element, NavigableString):
        return
    name = element.name
    if name in DOCX_HEADING_LEVELS:
        level = DOCX_HEADING_LEVELS[name]
        doc.add_heading(element.get_text(), level=level)
    elif name == "p":
        imgs = element.find_all("img", recursive=False)
        if len(imgs) == 1 and element.get_text().strip() == "":
            add_image(doc, imgs[0], base_dir)
            return
        for img in element.find_all("img"):
            add_image(doc, img, base_dir)
            img.decompose()
        if element.get_text().strip():
            add_paragraph(doc, element)
    elif name == "ul":
        for li in element.find_all("li", recursive=False):
            add_list_item(doc, li, ordered=False)
    elif name == "ol":
        for i, li in enumerate(element.find_all("li", recursive=False), 1):
            add_list_item(doc, li, ordered=True, index=i)
    elif name == "table":
        add_table(doc, element)
    elif name == "pre":
        code = element.find("code")
        text_content = code.get_text() if code else element.get_text()
        add_code_block(doc, text_content)
    elif name == "blockquote":
        add_blockquote(doc, element)
    elif name == "hr":
        doc.add_paragraph()
    elif name in ("div", "section"):
        for child in element.children:
            process_element(doc, child, base_dir)
    else:
        if element.get_text().strip():
            add_paragraph(doc, element)


def convert(md_path, docx_path, base_dir=None):
    if base_dir is None:
        base_dir = md_path.parent
    text = md_path.read_text(encoding="utf-8")
    html = md.markdown(text, extensions=["tables", "fenced_code", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
    base_style = doc.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    for child in soup.children:
        process_element(doc, child, base_dir)

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("md", type=Path)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--base-dir", type=Path, default=None)
    args = parser.parse_args()
    convert(args.md.resolve(), args.docx.resolve(),
            args.base_dir.resolve() if args.base_dir else None)
    return 0


if __name__ == "__main__":
    sys.exit(main())
