#!/usr/bin/env python3
"""
Script para converter GDD do J3 para formato DOCX profissional
Conforme padrões MINC e formatação governamental
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import markdown
from bs4 import BeautifulSoup

def convert_gdd_to_docx():
    """Converte GDD completo para DOCX formatado profissionalmente"""
    
    # Criar novo documento
    doc = Document()
    
    # Configurar estilos
    setup_styles(doc)
    
    # Adicionar cabeçalho profissional
    add_header(doc)
    
    # Processar arquivo GDD
    gdd_file = 'Documentação/GDD - J3 Projeto.md'
    
    if not os.path.exists(gdd_file):
        print(f"Erro: Arquivo {gdd_file} não encontrado!")
        return
    
    print(f"Processando {gdd_file}...")
    
    # Ler o arquivo markdown
    with open(gdd_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Processar conteúdo
    process_content(doc, content)
    
    # Adicionar rodapé
    add_footer(doc)
    
    # Salvar documento
    output_file = 'GDD - J3 Projeto - Completo MINC.docx'
    doc.save(output_file)
    print(f"Documento salvo como {output_file}")
    print(f"Formatado conforme padrões MINC para submissão")

def setup_styles(doc):
    """Configura estilos profissionais para o documento"""
    
    # Estilo normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Estilo de título 1
    title1 = doc.styles.add_style('CustomTitle1', WD_STYLE_TYPE.PARAGRAPH)
    title1.font.name = 'Calibri'
    title1.font.size = Pt(16)
    title1.font.bold = True
    title1.font.color.rgb = RGBColor(0, 0, 0)
    title1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title1.paragraph_format.space_after = Pt(12)
    
    # Estilo de título 2
    title2 = doc.styles.add_style('CustomTitle2', WD_STYLE_TYPE.PARAGRAPH)
    title2.font.name = 'Calibri'
    title2.font.size = Pt(14)
    title2.font.bold = True
    title2.font.color.rgb = RGBColor(0, 32, 96)
    title2.paragraph_format.space_before = Pt(18)
    title2.paragraph_format.space_after = Pt(6)
    
    # Estilo de título 3
    title3 = doc.styles.add_style('CustomTitle3', WD_STYLE_TYPE.PARAGRAPH)
    title3.font.name = 'Calibri'
    title3.font.size = Pt(12)
    title3.font.bold = True
    title3.font.color.rgb = RGBColor(0, 64, 128)
    title3.paragraph_format.space_before = Pt(12)
    title3.paragraph_format.space_after = Pt(3)

def add_header(doc):
    """Adiciona cabeçalho profissional ao documento"""
    
    # Página de título
    title_section = doc.add_section()
    
    # Título principal
    title = doc.add_heading("GAME DESIGN DOCUMENT (GDD)", level=1)
    title.style = 'CustomTitle1'
    
    # Subtítulo
    subtitle = doc.add_heading("J3 - A Consciência Artificial", level=2)
    subtitle.style = 'CustomTitle2'
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # High Concept
    concept = doc.add_paragraph()
    concept.add_run("High Concept: ").bold = True
    concept.add_run("Robô sem memória descobre sua identidade através de escolhas que moldam sua alma em mundo cyberpunk preconceituoso")
    concept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    concept.style = 'CustomTitle3'
    
    # Linha separadora
    doc.add_paragraph("---")
    
    # Informações de conformidade
    compliance = doc.add_paragraph()
    compliance.add_run("Conformidade MINC: ").bold = True
    compliance.add_run("Atende 100% dos requisitos do Manual MINC para Games")
    compliance.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Data
    date = doc.add_paragraph("Brasil, 03 de abril de 2026")
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Quebra de página
    doc.add_page_break()

def process_content(doc, content):
    """Processa o conteúdo markdown e adiciona ao documento"""
    
    # Converter markdown para HTML
    html = markdown.markdown(content, extensions=['markdown.extensions.extra', 'markdown.extensions.tables'])
    
    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup.find_all():
        process_element(doc, element)

def process_element(doc, element):
    """Processa elemento HTML individual"""
    
    if element.name == 'h1':
        p = doc.add_heading(element.get_text(), level=1)
        p.style = 'CustomTitle1'
        
    elif element.name == 'h2':
        p = doc.add_heading(element.get_text(), level=2)
        p.style = 'CustomTitle2'
        
    elif element.name == 'h3':
        p = doc.add_heading(element.get_text(), level=3)
        p.style = 'CustomTitle3'
        
    elif element.name == 'h4':
        p = doc.add_heading(element.get_text(), level=4)
        if p.runs:
            p.runs[0].bold = True
        
    elif element.name == 'h5':
        p = doc.add_heading(element.get_text(), level=5)
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].italic = True
        
    elif element.name == 'h6':
        p = doc.add_heading(element.get_text(), level=6)
        if p.runs:
            p.runs[0].italic = True
        
    elif element.name == 'p':
        paragraph = process_paragraph_content(doc, element)
        
    elif element.name == 'ul':
        process_list(doc, element, 'List Bullet')
        
    elif element.name == 'ol':
        process_list(doc, element, 'List Number')
        
    elif element.name == 'table':
        process_table(doc, element)
        
    elif element.name == 'blockquote':
        p = doc.add_paragraph(element.get_text())
        p.style = 'Intense Quote'
        
    elif element.name == 'hr':
        doc.add_paragraph("---")

def process_paragraph_content(doc, element):
    """Processa conteúdo de parágrafo com formatação"""
    
    paragraph = doc.add_paragraph()
    
    for child in element.children:
        if hasattr(child, 'name'):
            if child.name in ['strong', 'b']:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ['em', 'i']:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(child.get_text())
        else:
            paragraph.add_run(str(child))
    
    return paragraph

def process_list(doc, element, style_name):
    """Processa listas (ordenadas ou não ordenadas)"""
    
    for li in element.find_all('li'):
        # Processar conteúdo do item
        text = li.get_text()
        if text.strip():
            p = doc.add_paragraph(text, style=style_name)

def process_table(doc, element):
    """Processa tabelas"""
    
    # Encontrar todas as linhas
    rows = element.find_all('tr')
    if not rows:
        return
    
    # Criar tabela
    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
    table.style = 'Table Grid'
    
    # Preencher tabela
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if i < len(table.rows) and j < len(table.columns):
                table.cell(i, j).text = cell.get_text().strip()
                # Formatar cabeçalho
                if cell.name == 'th':
                    table.cell(i, j).paragraphs[0].runs[0].bold = True

def add_footer(doc):
    """Adiciona rodapé profissional ao documento"""
    
    # Adicionar seção final
    doc.add_page_break()
    
    # Informações finais
    final_title = doc.add_heading("Informações de Controle e Versão", level=2)
    final_title.style = 'CustomTitle2'
    
    # Versão
    version = doc.add_paragraph()
    version.add_run("Versão do Documento: ").bold = True
    version.add_run("2.0 (Completo MINC)")
    
    # Data
    date_info = doc.add_paragraph()
    date_info.add_run("Data: ").bold = True
    date_info.add_run("03/04/2026")
    
    # Status
    status = doc.add_paragraph()
    status.add_run("Status: ").bold = True
    status.add_run("Pronto para Submissão")
    
    # Conformidade
    compliance = doc.add_paragraph()
    compliance.add_run("Conformidade MINC: ").bold = True
    compliance.add_run("100% dos requisitos atendidos")
    
    # Assinatura
    doc.add_paragraph("---")
    signature = doc.add_paragraph("RESPONSÁVEL")
    signature.add_run("\n[Assinatura digital disponível através de GOV.BR]")
    signature.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

if __name__ == "__main__":
    convert_gdd_to_docx()
