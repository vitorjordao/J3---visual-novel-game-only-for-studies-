#!/usr/bin/env python3
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import markdown
from bs4 import BeautifulSoup

def convert_markdown_to_docx():
    # Create a new document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Files to process in order
    files_to_process = [
        'Roteiro/J3 - Roteiro.md',
        'Roteiro/Dias/Dia 1 - A Avenida.md',
        'Roteiro/Dias/Dia 2 - O Fliperama.md',
        'Roteiro/Dias/Dia 3 - O Beco.md',
        'Roteiro/Dias/Dia 4 - O Refúgio.md',
        'Roteiro/Dias/Dia 5 - O Cerco.md',
        'Roteiro/Dias/Dia 6 - A Revelação.md',
        'Roteiro/Dias/Dia 7 - O Final.md'
    ]
    
    for file_path in files_to_process:
        if not os.path.exists(file_path):
            print(f"Warning: File {file_path} not found, skipping...")
            continue
            
        print(f"Processing {file_path}...")
        
        # Add page break before each file except the first
        if file_path != files_to_process[0]:
            doc.add_page_break()
        
        # Read the markdown file
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Convert markdown to HTML
        html = markdown.markdown(content, extensions=['markdown.extensions.extra'])
        
        # Parse HTML and add to document
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all():
            if element.name == 'h1':
                p = doc.add_heading(element.get_text(), level=1)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.get_text(), level=4)
            elif element.name == 'h5':
                doc.add_heading(element.get_text(), level=5)
            elif element.name == 'h6':
                doc.add_heading(element.get_text(), level=6)
            elif element.name == 'p':
                # Handle bold and italic text
                paragraph = doc.add_paragraph()
                for child in element.children:
                    if hasattr(child, 'name'):
                        if child.name == 'strong' or child.name == 'b':
                            run = paragraph.add_run(child.get_text())
                            run.bold = True
                        elif child.name == 'em' or child.name == 'i':
                            run = paragraph.add_run(child.get_text())
                            run.italic = True
                        else:
                            paragraph.add_run(child.get_text())
                    else:
                        paragraph.add_run(str(child))
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    p = doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    p = doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'blockquote':
                p = doc.add_paragraph(element.get_text())
                p.style = 'Intense Quote'
            elif element.name == 'hr':
                doc.add_paragraph('---')
    
    # Save the document
    output_file = 'J3 - Roteiro Completo.docx'
    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    convert_markdown_to_docx()
